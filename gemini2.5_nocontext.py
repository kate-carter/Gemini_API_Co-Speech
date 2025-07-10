import google.generativeai as genai
import os
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import subprocess
import json
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import signal
import sys
import logging
from logging.handlers import RotatingFileHandler

# Configure logging based on user preference
def setup_logging(use_timestamped_logs):
    """
    Set up logging configuration based on user preference.
    
    Args:
        use_timestamped_logs (bool): Whether to use timestamped logs in terminal
    """
    # Create formatters
    detailed_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    simple_formatter = logging.Formatter('%(message)s')
    
    # Create handlers
    file_handler = RotatingFileHandler(
        'gemini_analysis.log',
        maxBytes=1024*1024,  # 1MB
        backupCount=5
    )
    file_handler.setFormatter(detailed_formatter)
    file_handler.setLevel(logging.DEBUG)
    
    console_handler = logging.StreamHandler()
    if use_timestamped_logs:
        console_handler.setFormatter(detailed_formatter)
    else:
        console_handler.setFormatter(simple_formatter)
    console_handler.setLevel(logging.INFO)
    
    # Configure root logger
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.DEBUG)
    root_logger.addHandler(file_handler)
    root_logger.addHandler(console_handler)

# --- Configuration ---
MODEL_NAME = "gemini-2.5-flash-preview-05-20"
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
OUTPUT_DIR = "/Users/Kate/Documents/CWRU/RedHen/GeminiOutput"
CONTEXT_PDF = "/Users/Kate/Documents/CWRU/RedHen/Relevant Studies/McNeill_CH3_PS.pdf"

# Analysis prompt
ANALYSIS_PROMPT = """Please analyze the co-speech gesture in this video in two sections:
1) The action being performed (visual analysis)
2) The co-speech gesture category the gesture belongs to (beat, deictic, iconic, etc.), 
with timestamps in MM:SS.MS for gesture onset.
If a gesture is iconic or metaphoric, please provide a description or subcategory."""

# Configure the model
generation_config = {
    "temperature": 0.4,  # Lower temperature for more focused responses
    "top_p": 0.95,      # Higher top_p for more reliable outputs
    "top_k": 32,        # Adjusted for vision model
    "max_output_tokens": 1024,  # Reduced for faster processing
    # "candidate_count": 1,  # Single candidate to avoid safety filtering issues - uncomment if needed
}

safety_settings = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

# Alternative safety settings format if still getting blocked:
# safety_settings = [
#     {
#         "category": "HARM_CATEGORY_HARASSMENT",
#         "threshold": "BLOCK_NONE"
#     },
#     {
#         "category": "HARM_CATEGORY_HATE_SPEECH", 
#         "threshold": "BLOCK_NONE"
#     },
#     {
#         "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
#         "threshold": "BLOCK_NONE"
#     },
#     {
#         "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
#         "threshold": "BLOCK_NONE"
#     },
#     {
#         "category": "HARM_CATEGORY_CIVIC_INTEGRITY",
#         "threshold": "BLOCK_NONE"
#     }
# ]

def format_text_with_bold(paragraph, text):
    """
    Format text with bold sections where text is surrounded by **asterisks** and
    italic sections where text is surrounded by *asterisks*.
    
    Args:
        paragraph: The paragraph object to add text to
        text: The text to format
    """
    try:
        # First split by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Remove asterisks and add as bold
                bold_text = part[2:-2]
                run = paragraph.add_run(bold_text)
                run.bold = True
            else:
                # Split the non-bold part by italic markers
                italic_parts = re.split(r'(\*.*?\*)', part)
                for italic_part in italic_parts:
                    if italic_part.startswith('*') and italic_part.endswith('*'):
                        # Remove asterisks and add as italic
                        italic_text = italic_part[1:-1]
                        run = paragraph.add_run(italic_text)
                        run.italic = True
                    else:
                        # Add as normal text
                        paragraph.add_run(italic_part)
    except Exception as e:
        logging.error(f"Error formatting text: {e}")
        # Fallback to plain text
        paragraph.add_run(text)

def format_time(seconds):
    """
    Format time in seconds to a human-readable string.
    
    Args:
        seconds (float): Time in seconds
        
    Returns:
        str: Formatted time string
    """
    try:
        if seconds < 60:
            return f"{seconds:.1f} seconds"
        else:
            minutes = int(seconds // 60)
            remaining_seconds = seconds % 60
            return f"{minutes} minutes and {remaining_seconds:.1f} seconds"
    except Exception as e:
        logging.error(f"Error formatting time: {e}")
        return "Unknown duration"

def get_video_duration(video_path):
    """
    Get the duration of a video file in seconds using ffprobe.
    
    Args:
        video_path (str): Path to the video file
        
    Returns:
        float: Duration in seconds, or None if there was an error
    """
    try:
        cmd = [
            'ffprobe',
            '-v', 'error',
            '-show_entries', 'format=duration',
            '-of', 'json',
            video_path
        ]
        logging.info(f"Running ffprobe command for: {video_path}")
        result = subprocess.run(cmd, capture_output=True, text=True)
        logging.debug(f"ffprobe output: {result.stdout}")
        
        if result.returncode == 0:
            data = json.loads(result.stdout)
            duration = float(data['format']['duration'])
            logging.info(f"Video duration: {format_time(duration)}")
            return duration
        else:
            logging.error(f"Error getting video duration: {result.stderr}")
            return None
    except Exception as e:
        logging.error(f"Error getting video duration: {e}")
        return None

def retry_failed_analysis(model, video_path, doc, video_filename, retry_count):
    """
    Retry analysis for a failed video file.
    
    Args:
        model: The Gemini model instance
        video_path: Path to the video file
        doc: The Word document
        video_filename: Name of the video file
        retry_count: Current retry attempt number
        
    Returns:
        tuple: (response_text, processing_time) if successful, (None, 0) if failed
    """
    logging.info(f"Retrying analysis for: {video_filename} (Attempt {retry_count})")
    uploaded_file_resource = None
    start_time = time.time()
    
    try:
        # Upload video
        logging.info(f"Uploading '{video_filename}'...")
        uploaded_file_resource = genai.upload_file(path=video_path,
                                                 display_name=video_filename)

        # Prepare prompt
        prompt_parts = [
            uploaded_file_resource,
            ANALYSIS_PROMPT
        ]

        # Get response
        logging.info("Analyzing gestures...")
        time.sleep(2)
        response = model.generate_content(prompt_parts)
        
        if response and hasattr(response, 'text'):
            response_text = response.text
            logging.info("Raw Gemini response (first 100 chars):")
            logging.info("-------------------")
            logging.info(response_text[:100] + "..." if len(response_text) > 100 else response_text)
            logging.info("-------------------")
            
            # Add retry response to document
            doc.add_paragraph(f'Retry Attempt {retry_count} for {video_filename}:', style='Heading 4')
            raw_para = doc.add_paragraph()
            format_text_with_bold(raw_para, response_text)
            
            # Add processing time
            end_time = time.time()
            processing_time = end_time - start_time
            doc.add_paragraph(f'Processing Time: {format_time(processing_time)}', style='Heading 4')
            
            return response_text, processing_time
        else:
            logging.warning(f"No text content in response for '{video_filename}'")
            return None, 0

    except Exception as e:
        logging.error(f"Error during retry for '{video_filename}': {e}")
        return None, 0

    finally:
        # Clean up uploaded file
        if uploaded_file_resource:
            try:
                time.sleep(1)
                genai.delete_file(uploaded_file_resource.name)
            except Exception as e:
                logging.error(f"Error deleting file during retry: {e}")

def create_analysis_document(video_paths):
    """
    Analyzes co-speech gestures and records the output in a Word document.
    
    Args:
        video_paths (list): A list of paths to video files.
    """
    if not GEMINI_API_KEY:
        logging.error("GEMINI_API_KEY environment variable not set.")
        return

    # Initialize request counter
    request_count = 0
    MAX_REQUESTS = 230
    
    doc = None  # Initialize doc variable at the start

    try:
        # Ask for initial permission for automated retries
        while True:
            auto_retry = input("\nDo you want to automate retries? (yes/no): ").lower()
            if auto_retry in ['yes', 'y', 'no', 'n']:
                break
            logging.warning("Please enter 'yes' or 'no'.")

        # Configure Gemini
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(
            model_name=MODEL_NAME,
            generation_config=generation_config,
            safety_settings=safety_settings
        )
        logging.info(f"Successfully initialized Gemini client with model: {MODEL_NAME}\n")

        # Create output directory if it doesn't exist
        os.makedirs(OUTPUT_DIR, exist_ok=True)

        # Create a new Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Co-Speech Gesture Analysis Report with Context (Gemini 2.5 Pro Vision)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add request count right after title
        doc.add_heading('Request Count', level=1)
        doc.add_paragraph(f'Total API requests made: {request_count}/{MAX_REQUESTS}')
        doc.add_paragraph()  # Single spacing after request count

        # Initialize tracking variables
        processing_times = []
        video_responses = {}  # Track responses for each video
        video_retry_counts = {}  # Track number of retries per video
        error_log = {}  # Track all failed attempts and their responses
        successful_processing_times = {}  # Track processing times for successful attempts only



        for video_path in video_paths:
            # Check if we've reached the request limit
            if request_count >= MAX_REQUESTS:
                logging.warning(f"Reached maximum request limit of {MAX_REQUESTS}. Saving current results...")
                break

            video_filename = os.path.basename(video_path)
            logging.info(f"Processing: {video_filename}")
            logging.info(f"Current request count: {request_count}/{MAX_REQUESTS}")
            logging.info("")  # Add blank line for readability

            if not os.path.exists(video_path):
                logging.error(f"Video file not found at '{video_path}'")
                error_log[video_filename] = ["File not found"]
                logging.info("")  # Add blank line after error
                continue

            # Start timing
            start_time = time.time()
            logging.info(f"Starting analysis at: {datetime.fromtimestamp(start_time).strftime('%H:%M:%S')}")

            # Get video duration
            video_duration = get_video_duration(video_path)
            duration_text = f" (Duration: {format_time(video_duration)})" if video_duration is not None else ""
            logging.info(f"Duration text for document: {duration_text}")
            
            uploaded_file_resource = None
            try:
                # Upload video
                logging.info(f"Uploading '{video_filename}'...")
                uploaded_file_resource = genai.upload_file(path=video_path,
                                                         display_name=video_filename)

                # Prepare prompt
                prompt_parts = [
                    uploaded_file_resource,
                    ANALYSIS_PROMPT
                ]

                # Get response
                logging.info("Analyzing gestures...")
                time.sleep(2)
                response = model.generate_content(prompt_parts)
                request_count += 1  # Increment request counter
                
                # Calculate end time right after getting the response
                end_time = time.time()
                processing_time = end_time - start_time
                
                if response and hasattr(response, 'text'):
                    response_text = response.text
                    logging.info("Raw Gemini response (first 100 chars):")
                    logging.info("-------------------")
                    logging.info(response_text[:100] + "..." if len(response_text) > 100 else response_text)
                    logging.info("-------------------")
                    
                    # Check if response is valid (has timestamps, sufficient length, and gesture types)
                    has_timestamps = bool(re.search(r'\d{2}:\d{2}\.\d{3}', response_text))
                    has_sufficient_length = len(response_text.split()) >= 50
                    has_gesture_types = any(token.lower() in response_text.lower() for token in ["beat", "metaphoric", "iconic", "deictic"])
                    
                    if not has_timestamps or not has_sufficient_length or not has_gesture_types:
                        error_log[video_filename] = [f"Initial attempt: {response_text}"]
                        logging.warning(f"Invalid response for {video_filename} - will retry (timestamps: {has_timestamps}, length: {has_sufficient_length}, gesture types: {has_gesture_types})")
                    else:
                        # Only store in video_responses if response is valid
                        video_responses[video_filename] = {
                            'response': response_text,
                            'duration': duration_text,
                            'processing_time': processing_time
                        }
                        video_retry_counts[video_filename] = 0
                        successful_processing_times[video_filename] = processing_time
                else:
                    error_log[video_filename] = ["No response received from Gemini"]

            except Exception as e:
                error_log[video_filename] = [f"Error analyzing video: {str(e)}"]
                logging.error(f"Error processing '{video_filename}': {e}")

            finally:
                # Clean up uploaded file
                if uploaded_file_resource:
                    try:
                        time.sleep(1)
                        genai.delete_file(uploaded_file_resource.name)
                    except Exception as e:
                        logging.error(f"Error deleting file: {e}")
                
                # Record processing time
                processing_times.append((video_filename, processing_time))
                logging.info(f"Processing completed at: {datetime.fromtimestamp(end_time).strftime('%H:%M:%S')}")
                logging.info(f"Total processing time: {format_time(processing_time)}")
                logging.info("")  # Add blank line after processing
                doc.add_paragraph()  # Add spacing

        # Process retries for files with errors
        retry_count = 0
        error_files = list(error_log.keys())
        while error_files and request_count < MAX_REQUESTS:
            logging.info(f"The following files still need valid output (Attempt {retry_count + 1}):")
            for error_file in error_files:
                logging.info(f"- {error_file}")
            logging.info("")  # Add blank line after listing error files
            
            # Only ask for permission if automated retries are not enabled
            if auto_retry not in ['yes', 'y']:
                while True:
                    retry_input = input(f"\nWould you like to retry analysis for these {len(error_files)} files? (yes/no): ").lower()
                    if retry_input in ['yes', 'y', 'no', 'n']:
                        break
                    logging.warning("Please enter 'yes' or 'no'.")
                
                if retry_input in ['no', 'n']:
                    logging.info("Stopping retry attempts.")
                    break
            else:
                logging.info("Automatically retrying analysis for failed files...")
            
            # Process retries
            for error_file in error_files:
                if request_count >= MAX_REQUESTS:
                    logging.warning(f"Reached maximum request limit of {MAX_REQUESTS}. Saving current results...")
                    break

                video_path = next((path for path in video_paths if os.path.basename(path) == error_file), None)
                if video_path and os.path.exists(video_path):
                    retry_response, retry_processing_time = retry_failed_analysis(model, video_path, doc, error_file, retry_count + 1)
                    request_count += 1  # Increment request counter
                    if retry_response:
                        has_timestamps = bool(re.search(r'\d{2}:\d{2}\.\d{3}', retry_response))
                        has_sufficient_length = len(retry_response.split()) >= 50
                        has_gesture_types = any(token.lower() in retry_response.lower() for token in ["beat", "metaphoric", "iconic", "deictic"])
                        
                        if has_timestamps and has_sufficient_length and has_gesture_types:
                            # Valid response - store in video_responses
                            video_duration = get_video_duration(video_path)
                            duration_text = f" (Duration: {format_time(video_duration)})" if video_duration is not None else ""
                            video_responses[error_file] = {
                                'response': retry_response,
                                'duration': duration_text,
                                'processing_time': retry_processing_time
                            }
                            video_retry_counts[error_file] = retry_count + 1
                            successful_processing_times[error_file] = retry_processing_time
                            logging.info(f"Successfully retried analysis for {error_file}")
                        else:
                            # Invalid response - add to error log
                            error_log[error_file].append(f"Retry attempt {retry_count + 1}: {retry_response}")
                            logging.warning(f"Retry failed for {error_file}")
                    else:
                        error_log[error_file].append(f"Retry attempt {retry_count + 1}: No response")
                        logging.warning(f"Retry failed for {error_file}")
            
            # Update error files list after retries
            error_files = [f for f in error_files if f not in video_responses]
            retry_count += 1
            
            if not error_files:
                logging.info("All files have been successfully processed!")
                break

        # Add request count to document
        doc.add_heading('Request Count', level=1)
        doc.add_paragraph(f'Total API requests made: {request_count}/{MAX_REQUESTS}')

        # Add successful analyses in alphabetical order
        if video_responses:
            doc.add_heading('Successful Analyses', level=1)
            for video_filename in sorted(video_responses.keys()):
                response_data = video_responses[video_filename]
                doc.add_heading(f'Video: {video_filename}{response_data["duration"]}', level=2)
                doc.add_paragraph('Analysis:', style='Heading 3')
                raw_para = doc.add_paragraph()
                format_text_with_bold(raw_para, response_data['response'])
                doc.add_paragraph('Processing Time:', style='Heading 3')
                doc.add_paragraph(f'Total time: {format_time(response_data["processing_time"])}')
                doc.add_paragraph()  # Single spacing between videos

        # Add error log section
        if error_log:  # Only add error log section if there are errors
            doc.add_heading('Error Log', level=1)
            # Sort error log entries by filename
            sorted_error_log = sorted(error_log.items(), key=lambda x: x[0])
            for filename, attempts in sorted_error_log:
                if filename not in video_responses:  # Only show files that never succeeded
                    doc.add_heading(f'File: {filename}', level=2)
                    for attempt in attempts:
                        doc.add_paragraph(attempt)
                    doc.add_paragraph()  # Single spacing between error entries

        # Add processing times summary
        doc.add_heading('Processing Times Summary', level=1)
        summary_table = doc.add_table(rows=1, cols=3)
        summary_table.style = 'Table Grid'
        
        # Add headers
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = 'Video'
        header_cells[1].text = 'Processing Time'
        header_cells[2].text = 'Retry Attempts'
        
        # Sort processing times by filename, but use successful processing times for display
        sorted_video_names = sorted(successful_processing_times.keys())
        
        # Add data rows
        for video_filename in sorted_video_names:
            processing_time = successful_processing_times[video_filename]
            row_cells = summary_table.add_row().cells
            row_cells[0].text = video_filename
            row_cells[1].text = format_time(processing_time)
            retries = video_retry_counts.get(video_filename, 0)
            row_cells[2].text = str(retries)
        
        # Add total processing time (only for successful attempts)
        total_time = sum(successful_processing_times.values())
        doc.add_paragraph(f'\nTotal processing time for all videos: {format_time(total_time)}')
        logging.info(f"Total processing time for all videos: {format_time(total_time)}")

    except Exception as e:
        logging.error(f"An error occurred during processing: {e}")
        raise  # Re-raise the exception to be caught by the outer try-finally

    finally:
        # Always ask about saving the document, even if there was an error
        if doc is not None:
            while True:
                save_doc = input("\nDo you want to save the analysis document? (yes/no): ").lower()
                if save_doc in ['yes', 'y', 'no', 'n']:
                    break
                logging.warning("Please enter 'yes' or 'no'.")

            if save_doc in ['yes', 'y']:
                # Save the document
                output_filename = f'pro2.5_NC_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}_requests_{request_count}.docx'
                output_path = os.path.join(OUTPUT_DIR, output_filename)
                doc.save(output_path)
                logging.info(f"Analysis saved to: {output_path}")
            else:
                logging.info("Analysis document not saved.")

def signal_handler(sig, frame):
    """
    Handle interruption signals (Ctrl+C) and ensure document save prompt.
    """
    logging.warning("Interrupted by user. Saving current results...")
    # Note: doc variable is not accessible in signal handler scope
    # This is a limitation of signal handlers in Python
    logging.warning("Cannot save document from signal handler - please use Ctrl+C to stop gracefully")
    sys.exit(0)

if __name__ == "__main__":
    # Set up signal handler for Ctrl+C
    signal.signal(signal.SIGINT, signal_handler)
    
    try:
        # Ask user about logging preference
        while True:
            log_pref = input("\nDo you want timestamped logs in the terminal? (yes/no): ").lower()
            if log_pref in ['yes', 'y', 'no', 'n']:
                break
            print("Please enter 'yes' or 'no'.")
        
        # Set up logging based on preference
        setup_logging(log_pref in ['yes', 'y'])
        
        logging.info("Co-Speech Gesture Analysis with Word Document Output (Gemini 2.5 Pro Vision)")
        logging.info("==================================================\n")
        
        # List of videos to analyze
        videos_to_analyze = [
        "/path/to/input/video1.mp4",
        "/path/to/input/video2.mp4",
        "/path/to/input/video3.mp4",
        ]

        if videos_to_analyze:
            create_analysis_document(videos_to_analyze)
        else:
            logging.warning("No video files specified for analysis.")
    except Exception as e:
        logging.error(f"An unexpected error occurred: {e}")
        sys.exit(1)
