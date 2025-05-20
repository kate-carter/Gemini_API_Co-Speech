import google.generativeai as genai
import os
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re

# --- Configuration ---
MODEL_NAME = "gemini-2.0-flash-001"
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")  # Set this in your environment
OUTPUT_DIR = "path/to/output/directory"  # Replace with your desired output directory

def format_text_with_bold(paragraph, text):
    """
    Format text with bold sections where text is surrounded by **asterisks**.
    
    Args:
        paragraph: The paragraph object to add text to
        text: The text to format
    """
    # Split the text by ** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Remove asterisks and add as bold
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            # Add as normal text
            paragraph.add_run(part)

def create_analysis_document(video_paths):
    """
    Analyzes co-speech gestures and records the output in a Word document.
    
    Args:
        video_paths (list): A list of paths to video files.
    """
    if not GEMINI_API_KEY:
        print("Error: GEMINI_API_KEY environment variable not set.")
        print("Please set it before running the script.")
        print("Example (in terminal): export GEMINI_API_KEY='YOUR_API_KEY'")
        return

    try:
        # Configure Gemini
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(MODEL_NAME)
        print(f"Successfully initialized Gemini client with model: {MODEL_NAME}\n")
    except Exception as e:
        print(f"Error initializing Gemini client: {e}")
        return

    # Create output directory if it doesn't exist
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Co-Speech Gesture Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Add spacing

    for video_path in video_paths:
        video_filename = os.path.basename(video_path)
        print(f"Processing: {video_filename}")

        if not os.path.exists(video_path):
            print(f"Error: Video file not found at '{video_path}'")
            continue

        # Add video section to document
        doc.add_heading(f'Video: {video_filename}', level=1)
        
        uploaded_file_resource = None
        try:
            # Upload video
            print(f"Uploading '{video_filename}'...")
            uploaded_file_resource = genai.upload_file(path=video_path,
                                                     display_name=video_filename)

            # Prepare prompt
            prompt_parts = [
                uploaded_file_resource,
                "Using linguistic terminology, can you classify the co-speech gesture in this clip in two sections: 1) the action being performed, and 2) the co speech gesture category the gesture belongs to?"
            ]

            # Get response
            print("Analyzing gestures...")
            time.sleep(2)
            response = model.generate_content(prompt_parts)

            if response.parts:
                # Parse and format the response
                response_text = response.text
                print(f"Raw Gemini response: {response_text}")  # Debug print
                
                # Add response to document
                doc.add_paragraph('Analysis:', style='Heading 2')
                
                # Try to split the response into action and category
                try:
                    # First, try to find the sections by looking for common patterns
                    if ("1)" in response_text and "2)" in response_text) or ("1." in response_text and "2." in response_text):
                        # Determine which format is being used
                        if "1)" in response_text:
                            # Split on "2)" and clean up the parts
                            parts = response_text.split("2)")
                            action = parts[0].replace("1)", "").strip()
                            category = "2)" + parts[1].strip()  # Keep the "2)" prefix for category
                        else:
                            # Split on "2." and clean up the parts
                            parts = response_text.split("2.")
                            action = parts[0].replace("1.", "").strip()
                            category = "2." + parts[1].strip()  # Keep the "2." prefix for category
                    elif "Action:" in response_text and "Category:" in response_text:
                        # Split on "Category:" and clean up the parts
                        parts = response_text.split("Category:")
                        action = parts[0].replace("Action:", "").strip()
                        category = "Category:" + parts[1].strip()  # Keep the "Category:" prefix
                    else:
                        # If no clear markers, try to split on newlines and look for meaningful content
                        lines = [line.strip() for line in response_text.split('\n') if line.strip()]
                        if len(lines) >= 2:
                            # Check if the second line starts with a number or category indicator
                            if lines[1].startswith(('2.', '2)', 'Category:', 'The gesture')):
                                action = lines[0]
                                category = lines[1]
                            else:
                                # If we can't determine the category line, use the whole response
                                action = response_text
                                category = "Category not clearly specified in response"
                        else:
                            # If we can't parse it, use the whole response
                            action = response_text
                            category = "Category not clearly specified in response"
                    
                    # Clean up any remaining numbering in the action if they weren't properly removed
                    action = action.replace("1)", "").replace("1.", "").strip()
                    
                    # Add formatted sections with clear headers
                    doc.add_paragraph('Action Performed:', style='Heading 3')
                    action_para = doc.add_paragraph()
                    format_text_with_bold(action_para, action)
                    
                    doc.add_paragraph('Co-Speech Gesture Category:', style='Heading 3')
                    category_para = doc.add_paragraph()
                    format_text_with_bold(category_para, category)
                    
                    # Add the raw response as reference
                    doc.add_paragraph('Raw Response for Reference:', style='Heading 3')
                    raw_para = doc.add_paragraph()
                    format_text_with_bold(raw_para, response_text)
                    
                except Exception as e:
                    # If parsing fails, add the raw response with error message
                    doc.add_paragraph('Error parsing response:', style='Heading 3')
                    doc.add_paragraph(f'Error details: {str(e)}')
                    doc.add_paragraph('Raw Response:', style='Heading 3')
                    raw_para = doc.add_paragraph()
                    format_text_with_bold(raw_para, response_text)
            else:
                doc.add_paragraph('No analysis received from Gemini.')

        except Exception as e:
            doc.add_paragraph(f'Error analyzing video: {str(e)}')
            print(f"Error processing '{video_filename}': {e}")

        finally:
            # Clean up uploaded file
            if uploaded_file_resource:
                try:
                    time.sleep(1)
                    genai.delete_file(uploaded_file_resource.name)
                except Exception as e:
                    print(f"Error deleting file: {e}")
            
            # Add spacing between videos
            doc.add_paragraph()

    # Save the document in the specified output directory
    output_filename = f'gesture_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    print(f"\nAnalysis saved to: {output_path}")

if __name__ == "__main__":
    print("Co-Speech Gesture Analysis with Word Document Output")
    print("==================================================\n")
    
    # List of videos to analyze - Replace these with your video paths
    videos_to_analyze = [
        "path/to/video1.mp4",
        "path/to/video2.mp4",
        "path/to/video3.mp4"
    ]

    if videos_to_analyze:
        create_analysis_document(videos_to_analyze)
    else:
        print("No video files specified for analysis.") 
