import google.generativeai as genai
import os
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import subprocess
import json
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import signal
import sys
import logging
from logging.handlers import RotatingFileHandler
import traceback

# --- Configuration ---
MODEL_NAME = "gemini-2.0-flash-001"
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
OUTPUT_DIR = "/path/to/output/directory"

# Configure the model
generation_config = {
    "temperature": 0.4,  # Lower temperature for more focused responses
    "top_p": 0.95,      # Higher top_p for more reliable outputs
    "top_k": 32,        # Adjusted for flash model
    "max_output_tokens": 1024,  # Reduced for faster processing
}

safety_settings = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

def format_text_with_bold(paragraph, text):
    """
    Format text with bold sections where text is surrounded by **asterisks** and
    italic sections where text is surrounded by *asterisks*.
    
    Args:
        paragraph: The paragraph object to add text to
        text: The text to format
    """
    # First split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Remove asterisks and add as bold
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            # Split the non-bold part by italic markers
            italic_parts = re.split(r'(\*.*?\*)', part)
            for italic_part in italic_parts:
                if italic_part.startswith('*') and italic_part.endswith('*'):
                    # Remove asterisks and add as italic
                    italic_text = italic_part[1:-1]
                    run = paragraph.add_run(italic_text)
                    run.italic = True
                else:
                    # Add as normal text
                    paragraph.add_run(italic_part)

def format_time(seconds):
    """
    Format time in seconds to a human-readable string.
    
    Args:
        seconds (float): Time in seconds
        
    Returns:
        str: Formatted time string
    """
    if seconds < 60:
        return f"{seconds:.1f} seconds"
    else:
        minutes = int(seconds // 60)
        remaining_seconds = seconds % 60
        return f"{minutes} minutes and {remaining_seconds:.1f} seconds"

def get_video_duration(video_path):
    """
    Get the duration of a video file in seconds using ffprobe.
    
    Args:
        video_path (str): Path to the video file
        
    Returns:
        float: Duration in seconds, or None if there was an error
    """
    try:
        cmd = [
            'ffprobe',
            '-v', 'error',
            '-show_entries', 'format=duration',
            '-of', 'json',
            video_path
        ]
        print(f"\nRunning ffprobe command for: {video_path}")
        result = subprocess.run(cmd, capture_output=True, text=True)
        print(f"ffprobe output: {result.stdout}")
        if result.returncode == 0:
            data = json.loads(result.stdout)
            duration = float(data['format']['duration'])
            print(f"Video duration: {format_time(duration)}")
            return duration
        else:
            print(f"Error getting video duration: {result.stderr}")
            return None
    except Exception as e:
        print(f"Error getting video duration: {e}")
        return None

def create_analysis_document(video_paths):
    """
    Analyzes co-speech gestures and records the output in a Word document.
    
    Args:
        video_paths (list): A list of paths to video files.
    """
    if not GEMINI_API_KEY:
        print("Error: GEMINI_API_KEY environment variable not set.")
        return

    try:
        # Configure Gemini
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(
            model_name=MODEL_NAME,
            generation_config=generation_config,
            safety_settings=safety_settings
        )
        print(f"Successfully initialized Gemini client with model: {MODEL_NAME}\n")
    except Exception as e:
        print(f"Error initializing Gemini client: {e}")
        return

    # Create output directory if it doesn't exist
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Co-Speech Gesture Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Add spacing

    # Add processing times summary section
    doc.add_heading('Processing Times Summary', level=1)
    processing_times = []

    for video_path in video_paths:
        video_filename = os.path.basename(video_path)
        print(f"\nProcessing: {video_filename}")

        if not os.path.exists(video_path):
            print(f"Error: Video file not found at '{video_path}'")
            continue

        # Start timing
        start_time = time.time()
        print(f"Starting analysis at: {datetime.fromtimestamp(start_time).strftime('%H:%M:%S')}")

        # Add video section to document
        video_duration = get_video_duration(video_path)
        duration_text = f" (Duration: {format_time(video_duration)})" if video_duration is not None else ""
        print(f"Duration text for document: {duration_text}")
        doc.add_heading(f'Video: {video_filename}{duration_text}', level=1)
        
        uploaded_file_resource = None
        try:
            # Upload video
            print(f"Uploading '{video_filename}'...")
            uploaded_file_resource = genai.upload_file(path=video_path,
                                                     display_name=video_filename)

            # Prepare prompt
            prompt_parts = [
                uploaded_file_resource,
                """Please analyze the co-speech gesture in this video in two sections:
1) The action being performed (visual analysis)
2) The co-speech gesture category the gesture belongs to (beat, deictic, iconic, etc.)"""
            ]

            # Get response
            print("Analyzing gestures...")
            time.sleep(2)
            response = model.generate_content(prompt_parts)

            if response.parts:
                # Parse and format the response
                response_text = response.text
                print(f"Raw Gemini response: {response_text}")  # Debug print
                
                # Add response to document
                doc.add_paragraph('Analysis:', style='Heading 2')
                
                # Try to split the response into action and category
                try:
                    # First, try to find the sections by looking for common patterns
                    if ("1)" in response_text and "2)" in response_text) or ("1." in response_text and "2." in response_text):
                        # Determine which format is being used
                        if "1)" in response_text:
                            # Split on "2)" and clean up the parts
                            parts = response_text.split("2)")
                            action = parts[0].replace("1)", "").strip()
                            category = "2)" + parts[1].strip()  # Keep the "2)" prefix for category
                        else:
                            # Split on "2." and clean up the parts
                            parts = response_text.split("2.")
                            action = parts[0].replace("1.", "").strip()
                            category = "2." + parts[1].strip()  # Keep the "2." prefix for category
                    elif "Action:" in response_text and "Category:" in response_text:
                        # Split on "Category:" and clean up the parts
                        parts = response_text.split("Category:")
                        action = parts[0].replace("Action:", "").strip()
                        category = "Category:" + parts[1].strip()  # Keep the "Category:" prefix
                    else:
                        # If no clear markers, try to split on newlines and look for meaningful content
                        lines = [line.strip() for line in response_text.split('\n') if line.strip()]
                        if len(lines) >= 2:
                            # Check if the second line starts with a number or category indicator
                            if lines[1].startswith(('2.', '2)', 'Category:', 'The gesture')):
                                action = lines[0]
                                category = lines[1]
                            else:
                                # If we can't determine the category line, use the whole response
                                action = response_text
                                category = "Category not clearly specified in response"
                        else:
                            # If we can't parse it, use the whole response
                            action = response_text
                            category = "Category not clearly specified in response"
                    
                    # Clean up any remaining numbering in the action if they weren't properly removed
                    action = action.replace("1)", "").replace("1.", "").strip()
                    
                    # Add formatted sections with clear headers
                    doc.add_paragraph('Action Performed:', style='Heading 3')
                    action_para = doc.add_paragraph()
                    format_text_with_bold(action_para, action)
                    
                    doc.add_paragraph('Co-Speech Gesture Category:', style='Heading 3')
                    category_para = doc.add_paragraph()
                    format_text_with_bold(category_para, category)
                    
                    # Add the raw response as reference
                    doc.add_paragraph('Raw Response for Reference:', style='Heading 3')
                    raw_para = doc.add_paragraph()
                    format_text_with_bold(raw_para, response_text)
                    
                except Exception as e:
                    # If parsing fails, add the raw response with error message
                    doc.add_paragraph('Error parsing response:', style='Heading 3')
                    doc.add_paragraph(f'Error details: {str(e)}')
                    doc.add_paragraph('Raw Response:', style='Heading 3')
                    raw_para = doc.add_paragraph()
                    format_text_with_bold(raw_para, response_text)
            else:
                doc.add_paragraph('No analysis received from Gemini.')

        except Exception as e:
            doc.add_paragraph(f'Error analyzing video: {str(e)}')
            print(f"Error processing '{video_filename}': {e}")

        finally:
            # Clean up uploaded file
            if uploaded_file_resource:
                try:
                    time.sleep(1)
                    genai.delete_file(uploaded_file_resource.name)
                except Exception as e:
                    print(f"Error deleting file: {e}")
            
            # Calculate and record processing time
            end_time = time.time()
            processing_time = end_time - start_time
            processing_times.append((video_filename, processing_time))
            print(f"Processing completed at: {datetime.fromtimestamp(end_time).strftime('%H:%M:%S')}")
            print(f"Total processing time: {format_time(processing_time)}")
            
            # Add processing time to the document
            doc.add_paragraph('Processing Time:', style='Heading 3')
            doc.add_paragraph(f'Total time: {format_time(processing_time)}')
            
            # Add spacing between videos
            doc.add_paragraph()

    # Add processing times summary at the beginning of the document
    doc.add_heading('Processing Times Summary', level=1)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    
    # Add headers
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = 'Video'
    header_cells[1].text = 'Processing Time'
    
    # Add data rows
    for video_filename, processing_time in processing_times:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = video_filename
        row_cells[1].text = format_time(processing_time)
    
    # Add total processing time
    total_time = sum(time for _, time in processing_times)
    doc.add_paragraph(f'\nTotal processing time for all videos: {format_time(total_time)}')
    print(f"\nTotal processing time for all videos: {format_time(total_time)}")

    # Save the document in the specified output directory
    output_filename = f'gesture_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    print(f"\nAnalysis saved to: {output_path}")

if __name__ == "__main__":
    print("Co-Speech Gesture Analysis with Word Document Output")
    print("==================================================\n")
    
    # List of videos to analyze
    videos_to_analyze = [
        "/path/to/input/video1.mp4",
        "/path/to/input/video2.mp4",
        "/path/to/input/video3.mp4"
    ]

    if videos_to_analyze:
        create_analysis_document(videos_to_analyze)
    else:
        print("No video files specified for analysis.") 

