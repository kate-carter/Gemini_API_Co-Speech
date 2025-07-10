import google.generativeai as genai
import os
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import subprocess
import json
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import signal
import sys
import logging
from logging.handlers import RotatingFileHandler

# Configure logging based on user preference
def setup_logging(use_timestamped_logs):
    """
    Set up logging configuration based on user preference.
    
    Args:
        use_timestamped_logs (bool): Whether to use timestamped logs in terminal
    """
    # Create formatters
    detailed_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    simple_formatter = logging.Formatter('%(message)s')
    
    # Create handlers
    file_handler = RotatingFileHandler(
        'gemini_analysis.log',
        maxBytes=1024*1024,  # 1MB
        backupCount=5
    )
    file_handler.setFormatter(detailed_formatter)
    file_handler.setLevel(logging.DEBUG)
    
    console_handler = logging.StreamHandler()
    if use_timestamped_logs:
        console_handler.setFormatter(detailed_formatter)
    else:
        console_handler.setFormatter(simple_formatter)
    console_handler.setLevel(logging.INFO)
    
    # Configure root logger
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.DEBUG)
    root_logger.addHandler(file_handler)
    root_logger.addHandler(console_handler)

# --- Configuration ---
MODEL_NAME = "gemini-2.0-flash-001"
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
OUTPUT_DIR = "/Users/Kate/Documents/CWRU/RedHen/GeminiOutput"
CONTEXT_PDF = "/Users/Kate/Documents/CWRU/RedHen/gemini_context/McNeill_CH3_PS.pdf"

# Analysis prompt
ANALYSIS_PROMPT = """Please analyze the co-speech gesture in this video in two sections:
1) The action being performed (visual analysis)
2) The co-speech gesture category the gesture belongs to, using the provided PDF as context. If the gesture is iconic or metaphoric, please provide a description or subcategory."""

# Configure the model
generation_config = {
    "temperature": 0.4,  # Lower temperature for more focused responses
    "top_p": 0.95,      # Higher top_p for more reliable outputs
    "top_k": 32,        # Adjusted for flash model
    "max_output_tokens": 1024,  # Reduced for faster processing
}

safety_settings = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

def format_text_with_bold(paragraph, text):
    """
    Format text with bold sections where text is surrounded by **asterisks** and
    italic sections where text is surrounded by *asterisks*.
    
    Args:
        paragraph: The paragraph object to add text to
        text: The text to format
    """
    try:
        # First split by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Remove asterisks and add as bold
                bold_text = part[2:-2]
                run = paragraph.add_run(bold_text)
                run.bold = True
            else:
                # Split the non-bold part by italic markers
                italic_parts = re.split(r'(\*.*?\*)', part)
                for italic_part in italic_parts:
                    if italic_part.startswith('*') and italic_part.endswith('*'):
                        # Remove asterisks and add as italic
                        italic_text = italic_part[1:-1]
                        run = paragraph.add_run(italic_text)
                        run.italic = True
                    else:
                        # Add as normal text
                        paragraph.add_run(italic_part)
    except Exception as e:
        logging.error(f"Error formatting text: {e}")
        # Fallback to plain text
        paragraph.add_run(text)

def format_time(seconds):
    """
    Format time in seconds to a human-readable string.
    
    Args:
        seconds (float): Time in seconds
        
    Returns:
        str: Formatted time string
    """
    try:
        if seconds < 60:
            return f"{seconds:.1f} seconds"
        else:
            minutes = int(seconds // 60)
            remaining_seconds = seconds % 60
            return f"{minutes} minutes and {remaining_seconds:.1f} seconds"
    except Exception as e:
        logging.error(f"Error formatting time: {e}")
        return "Unknown duration"

def get_video_duration(video_path):
    """
    Get the duration of a video file in seconds using ffprobe.
    
    Args:
        video_path (str): Path to the video file
        
    Returns:
        float: Duration in seconds, or None if there was an error
    """
    try:
        cmd = [
            'ffprobe',
            '-v', 'error',
            '-show_entries', 'format=duration',
            '-of', 'json',
            video_path
        ]
        logging.info(f"Running ffprobe command for: {video_path}")
        result = subprocess.run(cmd, capture_output=True, text=True)
        logging.debug(f"ffprobe output: {result.stdout}")
        
        if result.returncode == 0:
            data = json.loads(result.stdout)
            duration = float(data['format']['duration'])
            logging.info(f"Video duration: {format_time(duration)}")
            return duration
        else:
            logging.error(f"Error getting video duration: {result.stderr}")
            return None
    except Exception as e:
        logging.error(f"Error getting video duration: {e}")
        return None

def create_analysis_document(video_paths):
    """
    Analyzes co-speech gestures and records the output in a Word document.
    
    Args:
        video_paths (list): A list of paths to video files.
    """
    if not GEMINI_API_KEY:
        logging.error("GEMINI_API_KEY environment variable not set.")
        return

    # Initialize request counter and limit
    request_count = 0
    MAX_REQUESTS = 175
    doc = None  # Initialize doc variable at the start

    try:
        # Configure Gemini
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(
            model_name=MODEL_NAME,
            generation_config=generation_config,
            safety_settings=safety_settings
        )
        logging.info(f"Successfully initialized Gemini client with model: {MODEL_NAME}")
        logging.info(f"Model call limit set to: {MAX_REQUESTS}")
        logging.info("")  # Add blank line for readability

        # Create output directory if it doesn't exist
        os.makedirs(OUTPUT_DIR, exist_ok=True)

        # Create a new Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Gemini 2.0 Flash Analysis Report With Context', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add spacing

        # Add processing times summary section
        doc.add_heading('Processing Times Summary', level=1)
        processing_times = []

        # Upload context PDF if it exists
        context_resource = None
        if os.path.exists(CONTEXT_PDF):
            try:
                logging.info(f"Uploading context PDF: {CONTEXT_PDF}")
                context_resource = genai.upload_file(path=CONTEXT_PDF, display_name="context.pdf")
                logging.info("Context PDF uploaded successfully")
            except Exception as e:
                logging.error(f"Error uploading context PDF: {e}")

        for video_path in video_paths:
            # Check if we've reached the request limit
            if request_count >= MAX_REQUESTS:
                logging.warning(f"Reached maximum request limit of {MAX_REQUESTS}. Saving current results...")
                break

            video_filename = os.path.basename(video_path)
            logging.info(f"Processing: {video_filename}")
            logging.info(f"Current request count: {request_count}/{MAX_REQUESTS}")
            logging.info("")  # Add blank line for readability

            if not os.path.exists(video_path):
                logging.error(f"Video file not found at '{video_path}'")
                logging.info("")  # Add blank line after error
                continue

            # Start timing
            start_time = time.time()
            logging.info(f"Starting analysis at: {datetime.fromtimestamp(start_time).strftime('%H:%M:%S')}")

            # Add video section to document
            video_duration = get_video_duration(video_path)
            duration_text = f" (Duration: {format_time(video_duration)})" if video_duration is not None else ""
            logging.info(f"Duration text for document: {duration_text}")
            doc.add_heading(f'Video: {video_filename}{duration_text}', level=1)
            
            uploaded_file_resource = None
            processing_time = 0  # Initialize processing time
            try:
                # Upload video
                logging.info(f"Uploading '{video_filename}'...")
                uploaded_file_resource = genai.upload_file(path=video_path,
                                                         display_name=video_filename)

                # Prepare prompt with context
                prompt_parts = []
                
                # Add context PDF if available
                if context_resource:
                    prompt_parts.append(context_resource)
                    prompt_parts.append("Please use the above PDF as context for the following analysis.")
                
                # Add video and analysis request
                prompt_parts.extend([
                    uploaded_file_resource,
                    ANALYSIS_PROMPT
                ])

                # Get response
                logging.info("Analyzing gestures...")
                time.sleep(2)
                response = model.generate_content(prompt_parts)
                request_count += 1  # Increment request counter
                
                # Calculate processing time right after getting the response
                end_time = time.time()
                processing_time = end_time - start_time

                if response.parts:
                    # Parse and format the response
                    response_text = response.text
                    logging.info("Raw Gemini response (first 100 chars):")
                    logging.info("-------------------")
                    logging.info(response_text[:100] + "..." if len(response_text) > 100 else response_text)
                    logging.info("-------------------")
                    
                    # Add response to document
                    doc.add_paragraph('Analysis:', style='Heading 2')
                    
                    # Try to split the response into action and category
                    try:
                        # First, try to find the sections by looking for common patterns
                        if ("1)" in response_text and "2)" in response_text) or ("1." in response_text and "2." in response_text):
                            # Determine which format is being used
                            if "1)" in response_text:
                                # Split on "2)" and clean up the parts
                                parts = response_text.split("2)")
                                action = parts[0].replace("1)", "").strip()
                                category = "2)" + parts[1].strip()  # Keep the "2)" prefix for category
                            else:
                                # Split on "2." and clean up the parts
                                parts = response_text.split("2.")
                                action = parts[0].replace("1.", "").strip()
                                category = "2." + parts[1].strip()  # Keep the "2." prefix for category
                        elif "Action:" in response_text and "Category:" in response_text:
                            # Split on "Category:" and clean up the parts
                            parts = response_text.split("Category:")
                            action = parts[0].replace("Action:", "").strip()
                            category = "Category:" + parts[1].strip()  # Keep the "Category:" prefix
                        else:
                            # If no clear markers, try to split on newlines and look for meaningful content
                            lines = [line.strip() for line in response_text.split('\n') if line.strip()]
                            if len(lines) >= 2:
                                # Check if the second line starts with a number or category indicator
                                if lines[1].startswith(('2.', '2)', 'Category:', 'The gesture')):
                                    action = lines[0]
                                    category = lines[1]
                                else:
                                    # If we can't determine the category line, use the whole response
                                    action = response_text
                                    category = "Category not clearly specified in response"
                            else:
                                # If we can't parse it, use the whole response
                                action = response_text
                                category = "Category not clearly specified in response"
                        
                        # Clean up any remaining numbering in the action if they weren't properly removed
                        action = action.replace("1)", "").replace("1.", "").strip()
                        
                        # Add formatted sections with clear headers
                        doc.add_paragraph('Action Performed:', style='Heading 3')
                        action_para = doc.add_paragraph()
                        format_text_with_bold(action_para, action)
                        
                        doc.add_paragraph('Co-Speech Gesture Category:', style='Heading 3')
                        category_para = doc.add_paragraph()
                        format_text_with_bold(category_para, category)
                        
                        # Add processing time to document
                        doc.add_paragraph('Processing Time:', style='Heading 3')
                        doc.add_paragraph(f'Total time: {format_time(processing_time)}')
                        
                    except Exception as e:
                        logging.error(f"Error parsing response: {e}")
                        # If parsing fails, add the raw response
                        doc.add_paragraph('Raw Analysis:', style='Heading 3')
                        raw_para = doc.add_paragraph()
                        format_text_with_bold(raw_para, response_text)
                        
                        # Add processing time to document even for raw analysis
                        doc.add_paragraph('Processing Time:', style='Heading 3')
                        doc.add_paragraph(f'Total time: {format_time(processing_time)}')
                else:
                    logging.warning("No response received from Gemini")
                    doc.add_paragraph('No analysis received from Gemini.', style='Heading 3')
                    
                    # Add processing time to document even for failed analysis
                    doc.add_paragraph('Processing Time:', style='Heading 3')
                    doc.add_paragraph(f'Total time: {format_time(processing_time)}')

            except Exception as e:
                logging.error(f"Error processing video: {e}")
                doc.add_paragraph(f'Error analyzing video: {str(e)}', style='Heading 3')
                
                # Calculate processing time even for errors
                end_time = time.time()
                processing_time = end_time - start_time
                
                # Add processing time to document even for errors
                doc.add_paragraph('Processing Time:', style='Heading 3')
                doc.add_paragraph(f'Total time: {format_time(processing_time)}')

            finally:
                # Clean up uploaded file
                if uploaded_file_resource:
                    try:
                        time.sleep(1)
                        genai.delete_file(uploaded_file_resource.name)
                    except Exception as e:
                        logging.error(f"Error deleting file: {e}")
                
                # Record processing time for summary
                processing_times.append((video_filename, processing_time))
                logging.info(f"Processing completed at: {datetime.fromtimestamp(end_time).strftime('%H:%M:%S')}")
                logging.info(f"Processing time for '{video_filename}': {format_time(processing_time)}")
                logging.info("")  # Add blank line after processing
                doc.add_paragraph()  # Add spacing

        # Add request count to document
        doc.add_heading('Request Count', level=1)
        doc.add_paragraph(f'Total API requests made: {request_count}/{MAX_REQUESTS}')

        # Add processing times summary table
        if processing_times:
            summary_table = doc.add_table(rows=1, cols=2)
            summary_table.style = 'Table Grid'
            
            # Add headers
            header_cells = summary_table.rows[0].cells
            header_cells[0].text = 'Video'
            header_cells[1].text = 'Processing Time'
            
            # Sort processing times by filename
            sorted_processing_times = sorted(processing_times, key=lambda x: x[0])
            
            # Add data rows
            for video_filename, processing_time in sorted_processing_times:
                row_cells = summary_table.add_row().cells
                row_cells[0].text = video_filename
                row_cells[1].text = format_time(processing_time)

        # Add total processing time
        total_time = sum(time for _, time in processing_times)
        doc.add_paragraph(f'\nTotal processing time for all videos: {format_time(total_time)}')
        logging.info(f"Total processing time for all videos: {format_time(total_time)}")

    except Exception as e:
        logging.error(f"An error occurred during processing: {e}")
        raise  # Re-raise the exception to be caught by the outer try-finally

    finally:
        # Always ask about saving the document, even if there was an error
        if doc is not None:
            while True:
                save_doc = input("\nDo you want to save the analysis document? (yes/no): ").lower()
                if save_doc in ['yes', 'y', 'no', 'n']:
                    break
                logging.warning("Please enter 'yes' or 'no'.")

            if save_doc in ['yes', 'y']:
                # Save the document
                output_filename = f'2.0_context_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}_requests_{request_count}.docx'
                output_path = os.path.join(OUTPUT_DIR, output_filename)
                doc.save(output_path)
                logging.info(f"Analysis saved to: {output_path}")
            else:
                logging.info("Analysis document not saved.")

def signal_handler(sig, frame):
    """
    Handle interruption signals (Ctrl+C) and ensure document save prompt.
    """
    logging.warning("Interrupted by user. Saving current results...")
    # Note: doc variable is not accessible in signal handler scope
    # This is a limitation of signal handlers in Python
    logging.warning("Cannot save document from signal handler - please use Ctrl+C to stop gracefully")
    sys.exit(0)

if __name__ == "__main__":
    # Set up signal handler for Ctrl+C
    signal.signal(signal.SIGINT, signal_handler)
    
    try:
        # Ask user about logging preference
        while True:
            log_pref = input("\nDo you want timestamped logs in the terminal? (yes/no): ").lower()
            if log_pref in ['yes', 'y', 'no', 'n']:
                break
            print("Please enter 'yes' or 'no'.")
        
        # Set up logging based on preference
        setup_logging(log_pref in ['yes', 'y'])
        
        logging.info("Co-Speech Gesture Analysis with Word Document Output (Gemini 2.0 Flash)")
        logging.info("==================================================\n")
        
        # List of videos to analyze
        videos_to_analyze = [
        "/path/to/input/video1.mp4",
        "/path/to/input/video2.mp4",
        "/path/to/input/video3.mp4",
        ]

        if videos_to_analyze:
            create_analysis_document(videos_to_analyze)
        else:
            logging.warning("No video files specified for analysis.")
    except Exception as e:
        logging.error(f"An unexpected error occurred: {e}")
        sys.exit(1)
