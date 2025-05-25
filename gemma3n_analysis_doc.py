#note.... this script does not work because gemma 3n is not multimodal.  It was created for testing purposes to work with 
#the gemma API so if gemma gets multimodality it might be useful again (?) 
import os
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import subprocess
import json
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold

# --- Configuration ---
MODEL_NAME = "gemma-3n-e4b-it"  # Using Gemma 3n model
OUTPUT_DIR = "/path/to/output/directory"
MAX_TOKENS = 1400  # Maximum token limit

# Check for Google API key
if not os.getenv("GOOGLE_API_KEY"):
    print("Error: GOOGLE_API_KEY environment variable not set")
    print("Please set it using: export GOOGLE_API_KEY='your-api-key-here'")
    exit(1)

# Configure the API
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Configure the model
generation_config = {
    "temperature": 0.7,
    "top_p": 0.8,
    "top_k": 40,
    "max_output_tokens": MAX_TOKENS,
    "candidate_count": 1,  # Ensure we only get one response
}

safety_settings = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
}

def format_text_with_bold(paragraph, text):
    """
    Format text with bold sections where text is surrounded by **asterisks** and
    italic sections where text is surrounded by *asterisks*.
    
    Args:
        paragraph: The paragraph object to add text to
        text: The text to format
    """
    # First split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Remove asterisks and add as bold
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            # Split the non-bold part by italic markers
            italic_parts = re.split(r'(\*.*?\*)', part)
            for italic_part in italic_parts:
                if italic_part.startswith('*') and italic_part.endswith('*'):
                    # Remove asterisks and add as italic
                    italic_text = italic_part[1:-1]
                    run = paragraph.add_run(italic_text)
                    run.italic = True
                else:
                    # Add as normal text
                    paragraph.add_run(italic_part)

def format_time(seconds):
    """
    Format time in seconds to a human-readable string.
    
    Args:
        seconds (float): Time in seconds
        
    Returns:
        str: Formatted time string
    """
    if seconds < 60:
        return f"{seconds:.1f} seconds"
    else:
        minutes = int(seconds // 60)
        remaining_seconds = seconds % 60
        return f"{minutes} minutes and {remaining_seconds:.1f} seconds"

def get_video_duration(video_path):
    """
    Get the duration of a video file in seconds using ffprobe.
    
    Args:
        video_path (str): Path to the video file
        
    Returns:
        float: Duration in seconds, or None if there was an error
    """
    try:
        cmd = [
            'ffprobe',
            '-v', 'error',
            '-show_entries', 'format=duration',
            '-of', 'json',
            video_path
        ]
        print(f"\nRunning ffprobe command for: {video_path}")
        result = subprocess.run(cmd, capture_output=True, text=True)
        print(f"ffprobe output: {result.stdout}")
        if result.returncode == 0:
            data = json.loads(result.stdout)
            duration = float(data['format']['duration'])
            print(f"Video duration: {format_time(duration)}")
            return duration
        else:
            print(f"Error getting video duration: {result.stderr}")
            return None
    except Exception as e:
        print(f"Error getting video duration: {e}")
        return None

def create_analysis_document(video_paths):
    """
    Analyzes co-speech gestures and records the output in a Word document.
    
    Args:
        video_paths (list): A list of paths to video files.
    """
    try:
        # Initialize Gemma model
        print(f"Loading Gemma model: {MODEL_NAME}")
        model = genai.GenerativeModel(
            model_name=MODEL_NAME,
            generation_config=generation_config,
            safety_settings=safety_settings
        )
        print("Successfully initialized Gemma model\n")
    except Exception as e:
        print(f"Error initializing Gemma model: {e}")
        return

    # Create output directory if it doesn't exist
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Co-Speech Gesture Analysis Report (Gemma-3n-e4b-it)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Add spacing

    # Add processing times summary section
    doc.add_heading('Processing Times Summary', level=1)
    processing_times = []

    for video_path in video_paths:
        video_filename = os.path.basename(video_path)
        print(f"\nProcessing: {video_filename}")

        if not os.path.exists(video_path):
            print(f"Error: Video file not found at '{video_path}'")
            continue

        # Start timing
        start_time = time.time()
        print(f"Starting analysis at: {datetime.fromtimestamp(start_time).strftime('%H:%M:%S')}")

        # Get video duration
        video_duration = get_video_duration(video_path)
        duration_text = f" (Duration: {format_time(video_duration)})" if video_duration is not None else ""
        
        # Add video section to document
        doc.add_heading(f'Video: {video_filename}', level=1)
        doc.add_paragraph(f'Duration: {format_time(video_duration) if video_duration is not None else "Unknown"}')
        
        try:
            # Prepare prompt for Gemma
            prompt = """Please analyze the co-speech gesture in this video in two sections:
1) The action being performed
2) The co-speech gesture category the gesture belongs to (beat, deictic, iconic, etc.)"""

            # Generate response using Gemma with the video
            print(f"\nAttempting to analyze video: {video_filename}")
            print("Uploading video to Gemini...")
            try:
                # Verify file exists and is readable
                if not os.path.isfile(video_path):
                    raise FileNotFoundError(f"Video file not found: {video_path}")
                
                # Upload the video file first
                print(f"Uploading file: {video_path}")
                uploaded_file_resource = genai.upload_file(
                    path=video_path,
                    display_name=video_filename,
                    mime_type="video/mp4"  # Explicitly set MIME type
                )
                
                # Verify upload was successful
                if not uploaded_file_resource or not hasattr(uploaded_file_resource, 'name'):
                    raise Exception("File upload failed - no resource returned")
                
                print(f"Video uploaded successfully. Resource name: {uploaded_file_resource.name}")
                
                # Add a small delay to ensure the file is fully processed
                time.sleep(2)
                
                # Create prompt parts with the uploaded file
                prompt_parts = [uploaded_file_resource, prompt]
                
                print("Waiting for Gemini response...")
                response = model.generate_content(prompt_parts)
                
                if response and hasattr(response, 'text'):
                    response_text = response.text
                    print("\nRaw Gemini response (first 100 chars):")
                    print("-------------------")
                    print(response_text[:100] + "..." if len(response_text) > 100 else response_text)
                    print("-------------------")
                else:
                    print("Error: No response text received from Gemini")
                    print(f"Response object: {response}")
                    response_text = "No analysis received from Gemini."
                
                # Add raw response to document
                doc.add_paragraph('Analysis:')
                doc.add_paragraph(response_text)
                
            except Exception as upload_error:
                print(f"Error during video upload or analysis: {str(upload_error)}")
                print(f"Error type: {type(upload_error)}")
                if 'uploaded_file_resource' in locals():
                    try:
                        print("Attempting to clean up uploaded file...")
                        genai.delete_file(uploaded_file_resource.name)
                    except Exception as cleanup_error:
                        print(f"Error during cleanup: {cleanup_error}")
                raise upload_error
            finally:
                # Clean up uploaded file
                if 'uploaded_file_resource' in locals():
                    try:
                        time.sleep(1)  # Small delay to ensure processing is complete
                        genai.delete_file(uploaded_file_resource.name)
                        print("Successfully cleaned up uploaded file")
                    except Exception as e:
                        print(f"Error deleting uploaded file: {e}")
            
        except Exception as e:
            error_msg = f'Error analyzing video: {str(e)}'
            print(error_msg)
            doc.add_paragraph(error_msg)

        # Calculate and record processing time
        end_time = time.time()
        processing_time = end_time - start_time
        processing_times.append((video_filename, processing_time))
        print(f"Processing completed at: {datetime.fromtimestamp(end_time).strftime('%H:%M:%S')}")
        print(f"Total processing time: {format_time(processing_time)}")
        
        # Add processing time to the document
        doc.add_paragraph(f'Processing Time: {format_time(processing_time)}')
        
        # Add spacing between videos
        doc.add_paragraph()

    # Add processing times summary at the beginning of the document
    doc.add_heading('Processing Times Summary', level=1)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    
    # Add headers
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = 'Video'
    header_cells[1].text = 'Processing Time'
    
    # Add data rows
    for video_filename, processing_time in processing_times:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = video_filename
        row_cells[1].text = format_time(processing_time)
    
    # Add total processing time
    total_time = sum(time for _, time in processing_times)
    doc.add_paragraph(f'\nTotal processing time for all videos: {format_time(total_time)}')
    print(f"\nTotal processing time for all videos: {format_time(total_time)}")

    # Save the document in the specified output directory
    output_filename = f'gemma_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    print(f"\nAnalysis saved to: {output_path}")

if __name__ == "__main__":
    print("Co-Speech Gesture Analysis with Word Document Output (Gemma-3n-e4b-it)")
    print("==================================================\n")
    
    # List of videos to analyze
    videos_to_analyze = [
        "/path/to/video1.mp4",
        "/path/to/video2.mp4",
        "/path/to/video3.mp4",
    ]

    if videos_to_analyze:
        create_analysis_document(videos_to_analyze)
    else:
        print("No video files specified for analysis.") 
