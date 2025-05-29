import google.generativeai as genai
import os
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import subprocess
import json
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import signal
import sys

# --- Configuration ---
MODEL_NAME = "gemini-2.5-flash-preview-05-20"
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
OUTPUT_DIR = "/Users/Kate/Documents/CWRU/RedHen/GeminiOutput"

# Analysis prompt
ANALYSIS_PROMPT = """Please analyze the co-speech gesture in this video in two sections:
1) The action being performed (visual analysis)
2) The co-speech gesture category the gesture belongs to (beat, deictic, iconic, etc.), with timestamps in MM:SS.MS for gesture onset.
If a gesture is iconic or metaphoric, please provide a description or subcategory."""

# Configure the model
generation_config = {
    "temperature": 0.4,  # Lower temperature for more focused responses
    "top_p": 0.95,      # Higher top_p for more reliable outputs
    "top_k": 32,        # Adjusted for vision model
    "max_output_tokens": 1024,  # Reduced for faster processing
}

safety_settings = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

def format_text_with_bold(paragraph, text):
    """
    Format text with bold sections where text is surrounded by **asterisks** and
    italic sections where text is surrounded by *asterisks*.
    
    Args:
        paragraph: The paragraph object to add text to
        text: The text to format
    """
    # First split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Remove asterisks and add as bold
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            # Split the non-bold part by italic markers
            italic_parts = re.split(r'(\*.*?\*)', part)
            for italic_part in italic_parts:
                if italic_part.startswith('*') and italic_part.endswith('*'):
                    # Remove asterisks and add as italic
                    italic_text = italic_part[1:-1]
                    run = paragraph.add_run(italic_text)
                    run.italic = True
                else:
                    # Add as normal text
                    paragraph.add_run(italic_part)

def format_time(seconds):
    """
    Format time in seconds to a human-readable string.
    
    Args:
        seconds (float): Time in seconds
        
    Returns:
        str: Formatted time string
    """
    if seconds < 60:
        return f"{seconds:.1f} seconds"
    else:
        minutes = int(seconds // 60)
        remaining_seconds = seconds % 60
        return f"{minutes} minutes and {remaining_seconds:.1f} seconds"

def get_video_duration(video_path):
    """
    Get the duration of a video file in seconds using ffprobe.
    
    Args:
        video_path (str): Path to the video file
        
    Returns:
        float: Duration in seconds, or None if there was an error
    """
    try:
        cmd = [
            'ffprobe',
            '-v', 'error',
            '-show_entries', 'format=duration',
            '-of', 'json',
            video_path
        ]
        print(f"\nRunning ffprobe command for: {video_path}")
        result = subprocess.run(cmd, capture_output=True, text=True)
        print(f"ffprobe output: {result.stdout}")
        if result.returncode == 0:
            data = json.loads(result.stdout)
            duration = float(data['format']['duration'])
            print(f"Video duration: {format_time(duration)}")
            return duration
        else:
            print(f"Error getting video duration: {result.stderr}")
            return None
    except Exception as e:
        print(f"Error getting video duration: {e}")
        return None

def retry_failed_analysis(model, video_path, doc, video_filename, retry_count):
    """
    Retry analysis for a failed video file.
    
    Args:
        model: The Gemini model instance
        video_path: Path to the video file
        doc: The Word document
        video_filename: Name of the video file
        retry_count: Current retry attempt number
        
    Returns:
        str: The response text if successful, None if failed
    """
    print(f"\nRetrying analysis for: {video_filename} (Attempt {retry_count})")
    uploaded_file_resource = None
    start_time = time.time()
    try:
        # Upload video
        print(f"Uploading '{video_filename}'...")
        uploaded_file_resource = genai.upload_file(path=video_path,
                                                 display_name=video_filename)

        # Prepare prompt
        prompt_parts = [
            uploaded_file_resource,
            ANALYSIS_PROMPT
        ]

        # Get response
        print("Analyzing gestures...")
        time.sleep(2)
        response = model.generate_content(prompt_parts)
        
        if response and hasattr(response, 'text'):
            response_text = response.text
            print("\nRaw Gemini response (first 100 chars):")
            print("-------------------")
            print(response_text[:100] + "..." if len(response_text) > 100 else response_text)
            print("-------------------")
            
            # Add retry response to document
            doc.add_paragraph(f'Retry Attempt {retry_count} for {video_filename}:', style='Heading 4')
            raw_para = doc.add_paragraph()
            format_text_with_bold(raw_para, response_text)
            
            # Add processing time
            end_time = time.time()
            processing_time = end_time - start_time
            doc.add_paragraph(f'Processing Time: {format_time(processing_time)}', style='Heading 4')
            
            return response_text
        else:
            print("No analysis received from Gemini in retry attempt.")
            return None

    except Exception as e:
        print(f"Error during retry for '{video_filename}': {e}")
        return None

    finally:
        # Clean up uploaded file
        if uploaded_file_resource:
            try:
                time.sleep(1)
                genai.delete_file(uploaded_file_resource.name)
            except Exception as e:
                print(f"Error deleting file during retry: {e}")

def create_analysis_document(video_paths):
    """
    Analyzes co-speech gestures and records the output in a Word document.
    
    Args:
        video_paths (list): A list of paths to video files.
    """
    if not GEMINI_API_KEY:
        print("Error: GEMINI_API_KEY environment variable not set.")
        return

    # Initialize request counter
    request_count = 0
    MAX_REQUESTS = 450

    # Ask for initial permission for automated retries
    while True:
        auto_retry = input("\nDo you want to automate retries? (yes/no): ").lower()
        if auto_retry in ['yes', 'y', 'no', 'n']:
            break
        print("Please enter 'yes' or 'no'.")

    try:
        # Configure Gemini
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(
            model_name=MODEL_NAME,
            generation_config=generation_config,
            safety_settings=safety_settings
        )
        print(f"Successfully initialized Gemini client with model: {MODEL_NAME}\n")
    except Exception as e:
        print(f"Error initializing Gemini client: {e}")
        return

    # Create output directory if it doesn't exist
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Co-Speech Gesture Analysis Report (Gemini 2.5 Pro Vision)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Initialize tracking variables
    processing_times = []
    video_responses = {}  # Track responses for each video
    video_retry_counts = {}  # Track number of retries per video
    error_log = {}  # Track all failed attempts and their responses

    for video_path in video_paths:
        # Check if we've reached the request limit
        if request_count >= MAX_REQUESTS:
            print(f"\nReached maximum request limit of {MAX_REQUESTS}. Saving current results...")
            break

        video_filename = os.path.basename(video_path)
        print(f"\nProcessing: {video_filename}")
        print(f"Current request count: {request_count}/{MAX_REQUESTS}")

        if not os.path.exists(video_path):
            print(f"Error: Video file not found at '{video_path}'")
            error_log[video_filename] = ["File not found"]
            continue

        # Start timing
        start_time = time.time()
        print(f"Starting analysis at: {datetime.fromtimestamp(start_time).strftime('%H:%M:%S')}")

        # Get video duration
        video_duration = get_video_duration(video_path)
        duration_text = f" (Duration: {format_time(video_duration)})" if video_duration is not None else ""
        print(f"Duration text for document: {duration_text}")
        
        uploaded_file_resource = None
        try:
            # Upload video
            print(f"Uploading '{video_filename}'...")
            uploaded_file_resource = genai.upload_file(path=video_path,
                                                     display_name=video_filename)

            # Prepare prompt
            prompt_parts = [
                uploaded_file_resource,
                ANALYSIS_PROMPT
            ]

            # Get response
            print("Analyzing gestures...")
            time.sleep(2)
            response = model.generate_content(prompt_parts)
            request_count += 1  # Increment request counter
            
            # Calculate end time right after getting the response
            end_time = time.time()
            processing_time = end_time - start_time
            
            if response and hasattr(response, 'text'):
                response_text = response.text
                print("\nRaw Gemini response (first 100 chars):")
                print("-------------------")
                print(response_text[:100] + "..." if len(response_text) > 100 else response_text)
                print("-------------------")
                
                # Check if response is valid (has timestamps and sufficient length)
                if not re.search(r'\d{2}:\d{2}\.\d{3}', response_text) or len(response_text.split()) < 50:
                    error_log[video_filename] = [f"Initial attempt: {response_text}"]
                else:
                    # Only store in video_responses if response is valid
                    video_responses[video_filename] = {
                        'response': response_text,
                        'duration': duration_text,
                        'processing_time': processing_time
                    }
                    video_retry_counts[video_filename] = 0
            else:
                error_log[video_filename] = ["No response received from Gemini"]

        except Exception as e:
            error_log[video_filename] = [f"Error analyzing video: {str(e)}"]
            print(f"Error processing '{video_filename}': {e}")

        finally:
            # Clean up uploaded file
            if uploaded_file_resource:
                try:
                    time.sleep(1)
                    genai.delete_file(uploaded_file_resource.name)
                except Exception as e:
                    print(f"Error deleting file: {e}")
            
            # Record processing time
            processing_times.append((video_filename, processing_time))
            print(f"Processing completed at: {datetime.fromtimestamp(end_time).strftime('%H:%M:%S')}")
            print(f"Total processing time: {format_time(processing_time)}")
            doc.add_paragraph()  # Add spacing

    # Process retries for files with errors
    retry_count = 0
    error_files = list(error_log.keys())
    while error_files and request_count < MAX_REQUESTS:
        print(f"\nThe following files still need valid output (Attempt {retry_count + 1}):")
        for error_file in error_files:
            print(f"- {error_file}")
        
        # Only ask for permission if automated retries are not enabled
        if auto_retry not in ['yes', 'y']:
            while True:
                retry_input = input(f"\nWould you like to retry analysis for these {len(error_files)} files? (yes/no): ").lower()
                if retry_input in ['yes', 'y', 'no', 'n']:
                    break
                print("Please enter 'yes' or 'no'.")
            
            if retry_input in ['no', 'n']:
                print("Stopping retry attempts.")
                break
        else:
            print("\nAutomatically retrying analysis for failed files...")
        
        # Process retries
        for error_file in error_files:
            if request_count >= MAX_REQUESTS:
                print(f"\nReached maximum request limit of {MAX_REQUESTS}. Saving current results...")
                break

            video_path = next((path for path in video_paths if os.path.basename(path) == error_file), None)
            if video_path and os.path.exists(video_path):
                retry_response = retry_failed_analysis(model, video_path, doc, error_file, retry_count + 1)
                request_count += 1  # Increment request counter
                if retry_response:
                    if re.search(r'\d{2}:\d{2}\.\d{3}', retry_response) and len(retry_response.split()) >= 50:
                        # Valid response - store in video_responses
                        video_duration = get_video_duration(video_path)
                        duration_text = f" (Duration: {format_time(video_duration)})" if video_duration is not None else ""
                        video_responses[error_file] = {
                            'response': retry_response,
                            'duration': duration_text,
                            'processing_time': time.time() - start_time
                        }
                        video_retry_counts[error_file] = retry_count + 1
                        print(f"Successfully retried analysis for {error_file}")
                    else:
                        # Invalid response - add to error log
                        error_log[error_file].append(f"Retry attempt {retry_count + 1}: {retry_response}")
                        print(f"Retry failed for {error_file}")
                else:
                    error_log[error_file].append(f"Retry attempt {retry_count + 1}: No response")
                    print(f"Retry failed for {error_file}")
        
        # Update error files list after retries
        error_files = [f for f in error_files if f not in video_responses]
        retry_count += 1
        
        if not error_files:
            print("\nAll files have been successfully processed!")
            break

    # Add request count to document
    doc.add_heading('Request Count', level=1)
    doc.add_paragraph(f'Total API requests made: {request_count}/{MAX_REQUESTS}')

    # Add successful analyses in alphabetical order
    if video_responses:
        doc.add_heading('Successful Analyses', level=1)
        for video_filename in sorted(video_responses.keys()):
            response_data = video_responses[video_filename]
            doc.add_heading(f'Video: {video_filename}{response_data["duration"]}', level=2)
            doc.add_paragraph('Analysis:', style='Heading 3')
            raw_para = doc.add_paragraph()
            format_text_with_bold(raw_para, response_data['response'])
            doc.add_paragraph('Processing Time:', style='Heading 3')
            doc.add_paragraph(f'Total time: {format_time(response_data["processing_time"])}')
            doc.add_paragraph()  # Single spacing between videos

    # Add error log section
    if error_log:  # Only add error log section if there are errors
        doc.add_heading('Error Log', level=1)
        # Sort error log entries by filename
        sorted_error_log = sorted(error_log.items(), key=lambda x: x[0])
        for filename, attempts in sorted_error_log:
            if filename not in video_responses:  # Only show files that never succeeded
                doc.add_heading(f'File: {filename}', level=2)
                for attempt in attempts:
                    doc.add_paragraph(attempt)
                doc.add_paragraph()  # Single spacing between error entries

    # Add processing times summary
    doc.add_heading('Processing Times Summary', level=1)
    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = 'Table Grid'
    
    # Add headers
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = 'Video'
    header_cells[1].text = 'Processing Time'
    header_cells[2].text = 'Retry Attempts'
    
    # Sort processing times by filename
    sorted_processing_times = sorted(processing_times, key=lambda x: x[0])
    
    # Add data rows
    for video_filename, processing_time in sorted_processing_times:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = video_filename
        row_cells[1].text = format_time(processing_time)
        retries = video_retry_counts.get(video_filename, 0)
        row_cells[2].text = str(retries) if retries > 0 else "None"
    
    # Add total processing time
    total_time = sum(time for _, time in processing_times)
    doc.add_paragraph(f'\nTotal processing time for all videos: {format_time(total_time)}')
    print(f"\nTotal processing time for all videos: {format_time(total_time)}")

    # Save the document
    output_filename = f'2.5_520_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}_requests_{request_count}.docx'
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    print(f"\nAnalysis saved to: {output_path}")

def signal_handler(sig, frame):
    print("\n\nInterrupted by user. Saving current results...")
    if 'doc' in globals():
        # Save the document with current results
        output_filename = f'2.5_520_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}_interrupted.docx'
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        doc.save(output_path)
        print(f"\nAnalysis saved to: {output_path}")
    sys.exit(0)

if __name__ == "__main__":
    # Set up signal handler for Ctrl+C
    signal.signal(signal.SIGINT, signal_handler)
    
    print("Co-Speech Gesture Analysis with Word Document Output (Gemini 2.5 Pro Vision)")
    print("==================================================\n")
    
    # List of videos to analyze
    videos_to_analyze = [
        "/path/to/input/video1.mp4",
        "/path/to/input/video2.mp4",
        "/path/to/input/video3.mp4",
    ]

    if videos_to_analyze:
        create_analysis_document(videos_to_analyze)
    else:
        print("No video files specified for analysis.") 
